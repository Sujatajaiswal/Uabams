from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("UABAMS_Tomorrow_Meeting_Document_Updated.docx")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_table(table):
    table.style = "Table Grid"
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(9)
    for cell in table.rows[0].cells:
        set_cell_shading(cell, "E8EEF5")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True


def add_title(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("UABAMS Cloud Module Work Summary")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(11, 37, 69)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Meeting Presentation Document")
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(85, 85, 85)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Project: Unattended Axle Box Level Acceleration Measurement System (UABAMS)")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(85, 85, 85)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor(46, 116, 181) if level <= 2 else RGBColor(31, 77, 120)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_note_box(doc, title, text):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(31, 58, 95)
    p = cell.add_paragraph(text)
    p.paragraph_format.space_after = Pt(0)
    for run in p.runs:
        run.font.size = Pt(9)


def add_screenshot_slot(doc, caption):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F7F9FB")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(f"Insert Screenshot Here: {caption}")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(85, 85, 85)
    for _ in range(5):
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.1

    add_title(doc)
    add_note_box(
        doc,
        "One-line summary",
        "I implemented the cloud-side UABAMS workflow using FastAPI, Render PostgreSQL, and a PyQt5 desktop interface for calibration, threshold management, gateway data storage, and alert monitoring.",
    )

    add_heading(doc, "1. Work Completed", 1)
    add_bullets(
        doc,
        [
            "Built FastAPI backend APIs for wheel calibration, threshold management, gateway data ingestion, and alerts.",
            "Connected the backend to a Render PostgreSQL cloud database using SQLAlchemy and psycopg2.",
            "Built a PyQt5 desktop UI with five tabs: Wheel Calibration, Threshold, Gateway Data, Alerts, and Cloud Status.",
            "Added route-wise threshold support as required by the UABAMS specification.",
            "Added gateway packet metadata fields: record index, route name, KM marker, meter, status code, and 0.25 m sample distance.",
            "Implemented wheel wear compensation and corrected speed calculation.",
            "Implemented alert generation only when corrected speed is above 80 kmph and acceleration crosses the configured threshold.",
            "Added cloud readiness checks to show API status, Render PostgreSQL connection, schema readiness, and available database tables.",
            "Added a RailMAN export API so processed gateway and alert data can be handed over to RailMAN later.",
        ],
    )

    add_heading(doc, "2. Technology Stack", 1)
    table = doc.add_table(rows=1, cols=3)
    headers = ["Layer", "Technology", "Purpose"]
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, True)
    rows = [
        ("Desktop UI", "PyQt5", "Operator interface for forms and data tables."),
        ("Backend API", "Python + FastAPI", "Receives gateway data and exposes APIs."),
        ("Validation", "Pydantic", "Validates ranges such as +/-100g and threshold 0-100g."),
        ("Database", "PostgreSQL on Render", "Cloud storage for calibration, thresholds, gateway records, and alerts."),
        ("Database Access", "SQLAlchemy + psycopg2", "Executes PostgreSQL operations from FastAPI."),
        ("API Testing", "Swagger / OpenAPI", "Documents and tests the backend endpoints."),
        ("RailMAN Handoff", "FastAPI JSON Export", "Provides processed gateway and alert records for future RailMAN integration."),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    style_table(table)

    add_heading(doc, "3. System Architecture", 1)
    doc.add_paragraph("Current architecture:")
    add_numbered(
        doc,
        [
            "PyQt5 desktop app sends user input to FastAPI APIs.",
            "FastAPI validates and processes calibration, threshold, gateway, and alert data.",
            "FastAPI stores records in Render PostgreSQL cloud database.",
            "PyQt5 app refreshes data tables by calling GET APIs.",
            "Cloud Status verifies that the backend can reach the Render PostgreSQL database.",
            "RailMAN export provides a structured output from the cloud backend for future railway integration.",
        ],
    )
    add_note_box(
        doc,
        "Architecture flow",
        "Gateway Device -> FastAPI Backend -> Render PostgreSQL Cloud Database -> PyQt5 Desktop App / Web Dashboard / RailMAN Export",
    )

    add_heading(doc, "4. Application Workflow", 1)
    add_numbered(
        doc,
        [
            "Operator connects the PyQt5 app to the FastAPI URL.",
            "Operator enters wheel calibration data for train, axle, wheel position, wheel diameter, and encoder pulses.",
            "Operator enters route-wise vertical and lateral threshold values.",
            "Gateway sends spatial peak acceleration data with speed, KM, meter, GPS, and status.",
            "Backend stores the gateway record in PostgreSQL.",
            "Backend applies wheel wear correction to calculate corrected speed.",
            "Backend checks corrected speed above 80 kmph and compares acceleration with route thresholds.",
            "If limit is crossed, backend stores an alert record.",
            "PyQt5 UI displays stored gateway data and generated alerts.",
            "Cloud Status page confirms that the API and cloud database are ready.",
            "RailMAN export preview shows the processed data format that can be consumed by RailMAN later.",
        ],
    )

    add_heading(doc, "5. Screen-by-Screen UI Explanation", 1)
    screens = [
        (
            "Page 1: Wheel Calibration",
            "This page stores wheel wear calibration values. It captures train number, axle number, wheel position, new wheel diameter, current wheel diameter, and encoder pulses per revolution. The backend calculates wheel wear, circumference, distance per pulse, and correction factor.",
            "While presenting, explain that this satisfies the wheel wear compensation requirement for accurate speed and distance measurement.",
        ),
        (
            "Page 2: Threshold Management",
            "This page stores route-wise vertical and lateral acceleration thresholds. Threshold values are restricted from 0g to 100g.",
            "While presenting, explain that these are editable predefined limits used for alert generation.",
        ),
        (
            "Page 3: Gateway Data",
            "This page simulates/receives gateway data. It stores record index, train number, route, KM marker, meter, vertical/lateral acceleration, speed, GPS location, and status code.",
            "While presenting, explain that the gateway data is stored in cloud PostgreSQL and includes 0.25 m spatial interval metadata.",
        ),
        (
            "Page 4: Alerts",
            "This page displays generated alerts. Each alert includes train, route, alert type, measured value, threshold value, corrected speed, KM/meter, status, location, and timestamp.",
            "While presenting, explain that alerts are generated only above 80 kmph and when acceleration exceeds configured limits.",
        ),
        (
            "Page 5: Cloud Status",
            "This page verifies the cloud side of the project. It shows API status, database status, Render PostgreSQL host, schema readiness, available tables, gateway ingest endpoint, and RailMAN export endpoint.",
            "While presenting, explain that the database is already cloud-based on Render PostgreSQL. The local FastAPI server is the development backend and can later be deployed on Render/RailMAN infrastructure.",
        ),
    ]
    for title, explanation, talking_point in screens:
        add_heading(doc, title, 2)
        doc.add_paragraph(explanation)
        add_note_box(doc, "Talking point", talking_point)
        add_screenshot_slot(doc, title)

    add_heading(doc, "6. Calculations Implemented", 1)
    calc_table = doc.add_table(rows=1, cols=3)
    for i, h in enumerate(["Calculation", "Formula", "Example"]):
        set_cell_text(calc_table.rows[0].cells[i], h, True)
    calculations = [
        ("Wheel wear", "new_diameter - current_diameter", "920 - 890 = 30 mm"),
        ("Wheel circumference", "pi * current_diameter", "3.14159 * 890 = 2796.01 mm"),
        ("Distance per pulse", "circumference / encoder_pulses", "2796.01 / 1024 = 2.73 mm"),
        ("Correction factor", "current_diameter / new_diameter", "890 / 920 = 0.9674"),
        ("Corrected speed", "gateway_speed * correction_factor", "90 * 0.9674 = 87.06 kmph"),
        ("Vertical alert", "corrected_speed > 80 AND abs(vertical_g) > vertical_threshold", "87.06 > 80 and 75 > 50"),
        ("Lateral alert", "corrected_speed > 80 AND abs(lateral_g) > lateral_threshold", "Only triggers if lateral value crosses limit"),
    ]
    for row in calculations:
        cells = calc_table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    style_table(calc_table)

    add_heading(doc, "7. Database Work", 1)
    doc.add_paragraph("The backend creates or updates the following database tables automatically:")
    db_table = doc.add_table(rows=1, cols=2)
    set_cell_text(db_table.rows[0].cells[0], "Table", True)
    set_cell_text(db_table.rows[0].cells[1], "Purpose", True)
    db_rows = [
        ("wheel_calibration", "Stores train/axle/wheel calibration, wheel wear, correction factor, and distance per pulse."),
        ("thresholds", "Stores route-wise vertical and lateral acceleration thresholds."),
        ("acceleration_data", "Stores gateway data, corrected speed, route/KM/meter/status, GPS, and sample distance."),
        ("alerts", "Stores generated vertical/lateral alerts with measured value, threshold, corrected speed, route/KM/meter, and GPS."),
    ]
    for row in db_rows:
        cells = db_table.add_row().cells
        set_cell_text(cells[0], row[0])
        set_cell_text(cells[1], row[1])
    style_table(db_table)

    add_heading(doc, "8. Cloud And RailMAN Readiness", 1)
    add_note_box(
        doc,
        "Current cloud status",
        "The PostgreSQL database is hosted on Render, so calibration, threshold, gateway, and alert data are stored in the cloud. During development, FastAPI still runs locally at http://127.0.0.1:8000. For production, FastAPI should be deployed on Render or the final railway/RailMAN-approved server.",
    )
    add_bullets(
        doc,
        [
            "Cloud database: Render PostgreSQL is used through DATABASE_URL.",
            "Cloud proof endpoint: /cloud/status shows database connection, schema readiness, available tables, and API health.",
            "Gateway cloud flow: gateway data is posted to /api/data and stored in PostgreSQL.",
            "RailMAN handoff: /railman/export returns latest gateway and alert records in a structured JSON response.",
            "Future production step: deploy the FastAPI backend to a public cloud URL so gateway/RailMAN systems can access it directly.",
        ],
    )

    add_heading(doc, "9. API Endpoints Implemented", 1)
    api_table = doc.add_table(rows=1, cols=3)
    for i, h in enumerate(["Endpoint", "Method", "Purpose"]):
        set_cell_text(api_table.rows[0].cells[i], h, True)
    api_rows = [
        ("/wheel-calibration", "POST/GET", "Save and view wheel calibration records."),
        ("/threshold", "POST/GET", "Save and view route-wise threshold values."),
        ("/api/data", "POST/GET", "Receive and view gateway data stored in cloud."),
        ("/alerts", "GET", "View generated alerts."),
        ("/cloud/status", "GET", "Verify API status, Render PostgreSQL connection, schema readiness, and RailMAN readiness."),
        ("/railman/export", "GET", "Export processed gateway and alert records for future RailMAN integration."),
        ("/docs", "GET", "Swagger API documentation."),
    ]
    for row in api_rows:
        cells = api_table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    style_table(api_table)

    add_heading(doc, "10. Specification Compliance Mapping", 1)
    spec_table = doc.add_table(rows=1, cols=3)
    for i, h in enumerate(["Specification Requirement", "Implemented Work", "Status"]):
        set_cell_text(spec_table.rows[0].cells[i], h, True)
    spec_rows = [
        ("Axle acceleration up to +/-100g", "Backend validates vertical and lateral values between -100g and +100g.", "Completed"),
        ("Editable predefined limits", "Route-wise vertical and lateral threshold management added.", "Completed"),
        ("Alerts above 80 kmph", "Alert logic checks corrected_speed_kmph > 80 before generating alert.", "Completed"),
        ("Wheel wear compensation", "Correction factor and corrected speed are calculated from wheel diameters.", "Completed"),
        ("Spatial acceleration interval 25 cm", "Gateway records store sample_distance_m = 0.25.", "Completed"),
        ("Data stored in database", "Render PostgreSQL stores calibration, threshold, gateway, and alert data.", "Completed"),
        ("Gateway packet details", "Record index, speed, KM, meter, status, latitude, and longitude are stored with gateway records.", "Completed"),
        ("Railway/RailMAN future use", "RailMAN export endpoint is prepared for structured cloud data handoff.", "Prototype ready"),
    ]
    for row in spec_rows:
        cells = spec_table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    style_table(spec_table)

    add_heading(doc, "11. Demo Workflow For Tomorrow", 1)
    add_numbered(
        doc,
        [
            "Start FastAPI backend using uvicorn.",
            "Open the PyQt5 desktop application.",
            "Click Connect and confirm the API URL is http://127.0.0.1:8000 for local demo.",
            "Open Cloud Status and click Refresh Cloud Status to show database connection and schema readiness.",
            "Open Threshold and set route DEFAULT with vertical/lateral threshold values.",
            "Open Wheel Calibration and save calibration values for train, axle, and wheel.",
            "Open Gateway Data and send sample data with speed above 80 kmph.",
            "Open Alerts and show the generated alert record.",
            "Return to Cloud Status and preview RailMAN Export to show the future RailMAN handoff data.",
        ],
    )

    add_heading(doc, "12. 500 Error Explanation And Demo Fix", 1)
    add_note_box(
        doc,
        "Observed error",
        "The PyQt app showed 500 Internal Server Error for /threshold?route_name=DEFAULT. This means the desktop UI reached FastAPI, but the backend endpoint failed while reading threshold data.",
    )
    add_bullets(
        doc,
        [
            "Most likely reason during development: the backend server was running old code or an older database schema before the latest changes were loaded.",
            "Fix before demo: stop uvicorn using Ctrl+C, restart FastAPI with --reload, restart the PyQt app, and click Connect again.",
            "Verification: /threshold, /cloud/status, and /railman/export were checked after the update and are working.",
        ],
    )

    add_heading(doc, "13. Meeting Script", 1)
    add_bullets(
        doc,
        [
            "I implemented the cloud-side UABAMS module using FastAPI and Render PostgreSQL.",
            "I added a PyQt5 desktop interface for operators to manage calibration, thresholds, gateway data, alerts, and cloud status.",
            "The wheel calibration page calculates wheel wear and correction factor for accurate speed and distance.",
            "The threshold page supports route-wise vertical and lateral limits from 0g to 100g.",
            "The gateway page stores spatial gateway data with index, KM, meter, GPS, speed, and status.",
            "The alert page shows alerts generated only when corrected speed is above 80 kmph and acceleration crosses threshold.",
            "The Cloud Status page proves that the API can connect to the Render PostgreSQL cloud database.",
            "The RailMAN export endpoint is prepared so later RailMAN can consume processed gateway and alert records.",
            "The database schema is aligned with the specification and supports cloud storage for reports/data.",
        ],
    )

    add_heading(doc, "14. Future Enhancements", 1)
    add_bullets(
        doc,
        [
            "Deploy the FastAPI backend itself on Render so gateway devices can post directly to a public cloud URL.",
            "Add authentication for operator login and protected API access.",
            "Add export options for CSV/PDF reports.",
            "Add SMS/notification integration for generated alerts.",
            "Add 50 m highest-peak grouping logic if required for final alert notification compliance.",
            "Add dashboard charts for speed, vertical acceleration, lateral acceleration, and alert trend.",
        ],
    )

    section = doc.sections[0]
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("UABAMS Cloud Module Work Summary").font.size = Pt(9)

    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
